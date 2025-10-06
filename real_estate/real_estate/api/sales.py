import frappe
import click
from frappe import _


from real_estate.real_estate.data_engine.query import OfferQuery
from real_estate.real_estate.data_engine.query import SalesQuery
from real_estate.real_estate.data_engine.query import PaymentsQuery
from erpnext.selling.doctype.sales_order.sales_order import make_sales_invoice
from frappe.utils.weasyprint import download_pdf
from frappe.www.printview import validate_print_permission
from frappe.translate import print_language
from io import BytesIO
from pdf2docx import Converter
import tempfile
from docx import Document  # Import python-docx for Word editing
from docx.shared import Pt
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

@frappe.whitelist()
def create_offer(plot, project, customer, amount, date, discount):

	user = frappe.session.user

	employee = frappe.get_all("Employee", filters=[["user_id", "=", user]], fields=["name", "employee_name", "branch"])[0]

	sales = frappe.new_doc("Sales Order")
	sales.customer = customer
	sales.sales_rep = employee.name
	sales.project = project
	sales.branch = employee.branch
	sales.plot = plot
	sales.discount_amount = discount
	item = sales.append("items",{})

	# has_serial = frappe.get_value("Item", "Plot", has_serial_)

	item.item_code = 'Plot'
	item.item_name = plot
	item.qty = 1
	item.rate = amount
	sales.transaction_date = date
	sales.set_missing_values()

	sales.insert()
	hash = frappe.generate_hash(length=56) + sales.name.split("-")[-1]
	frappe.db.set_value("Sales Order", sales.name, "custom_link_code", hash)
	frappe.db.set_value("Sales Order", sales.name, "custom_offer__letter_template_url", 
						get_short_link(frappe.utils.get_url() + "/api/method/real_estate.real_estate.api.downloads.download_offer_template?ref=" + hash))
	frappe.db.set_value("Sales Order", sales.name, "custom_signed_offer_letter_url", 
						get_short_link(frappe.utils.get_url() + "/api/method/real_estate.real_estate.api.downloads.download_signed_offer?ref=" + hash))
	frappe.db.set_value("Sales Order", sales.name, "custom_agreement_template_url", 
						get_short_link(frappe.utils.get_url() + "/api/method/real_estate.real_estate.api.downloads.download_agreement_template?ref=" + hash))
	frappe.db.set_value("Sales Order", sales.name, "custom_signed_agreement_url", 
						get_short_link(frappe.utils.get_url() + "/api/method/real_estate.real_estate.api.downloads.download_signed_agreement?ref=" + hash))
	frappe.db.set_value("Sales Order", sales.name, "custom_statement_url", 
						get_short_link(frappe.utils.get_url() + "/api/method/real_estate.real_estate.api.downloads.download_statement?ref=" + hash))

	return sales.name

def get_short_link(long_url):
	import requests
	import json

	# Define the endpoint and authentication token
	url = 'https://api.tinyurl.com/create'
	token = 'OftLz7Wt0rr93wS6HuoGDjD3g7aUnI68tHdcIi4PJ5UEXDobdPpR21Q5YTWP'  # Replace with your actual token

	# Prepare the JSON data
	data = {
		"url": long_url,
		"domain": "tinyurl.com",
		"description": "string"
	}

	# Prepare headers with bearer authentication
	headers = {
		'Authorization': f'Bearer {token}',
		'Content-Type': 'application/json'
	}

	response = requests.post(url, headers=headers, data=json.dumps(data))

	# Check if the request was successful
	if response.status_code == 200:
		# Parse the response JSON
		response_data = response.json()
		
		# Access the relevant data from the response
		tiny_url = response_data['data']['tiny_url']
		
		return tiny_url
	else:
		return long_url,

@frappe.whitelist()
def get_html_agreement(doc_name):
	doc = frappe.get_doc("Sales Order", doc_name)
	if doc.agreement is not None and doc.agreement != '':
		return doc.agreement

	generator = PrintFormatGenerator(doc.agreement_template, doc, "NO letter head")
	return generator.get_main_html()
	

@frappe.whitelist()
def save_agreement(doc_name, agreement):
	frappe.db.set_value("Sales Order", doc_name, "agreement", agreement)

@frappe.whitelist()
def get_offer(offer):
	offer = frappe.get_doc("Sales Order", offer)
	customer = frappe.get_doc("Customer", offer.customer)
	image = ""
	if customer.image:
		image = frappe.utils.get_url() + "/" + customer.image

	sale_state = offer.sale_state
	if offer.docstatus == 2:
		sale_state = "Cancelled"

	return {
		"name": offer.name,
		"customer": customer.name,
		"plot": offer.plot,
		"project": offer.project,
		"sale_state": sale_state,
		"agreement_template": offer.agreement_template,
		"include_lawyer": offer.include_lawyer,
		"total": offer.total,
		"net_total": offer.net_total,
		"advance_paid": offer.advance_paid,
		"sales_rep": offer.sales_rep,
		"discount_amount": offer.discount_amount,
		"transaction_date": offer.transaction_date,
		"status": offer.status,
		"payment_terms": offer.payment_schedule,
		"files": offer.files,
		"currency": offer.currency,
		"custom_purchase_method": offer.custom_purchase_method,
		"custom_booking_fee": offer.custom_booking_fee,
		"custom_discount_approved": offer.custom_discount_approved,
		"custom_installments_approved": offer.custom_installments_approved
	}

@frappe.whitelist()
def get_sale(sale):
	sale = frappe.get_doc("Sales Invoice", sale)
	customer = frappe.get_doc("Customer", sale.customer)
	image = ""
	if customer.image:
		image = frappe.utils.get_url() + "/" + customer.image

	return {
		"name": sale.name,
		"customer": {'name': customer.name, 'image': image},
		"plot": sale.plot,
		"transaction_date": sale.posting_date,
		"status": sale.status,
		"due_date": sale.due_date,
		"branch": sale.branch,
		"sales_rep": sale.sales_rep,
		"net_total": sale.net_total,
		"status": sale.status,
		"payment_terms": sale.payment_schedule,
		"files": sale.files
	}


def update_plots_on_offer_submit(doc, method=None):
	project = frappe.get_doc("Project", doc.project)
	for pl in project.plots:
		if pl.name == doc.plot:
			if pl.status != "ON HOLD":
				frappe.throw("The plot in the offer is NOT ON HOLD")

	frappe.db.set_value("Project Plot Detail", doc.plot, "status", "BOOKED")

def update_plots_on_offer_cancel(doc, method=None):
	frappe.db.set_value("Project Plot Detail", doc.plot, "customer", None)
	frappe.db.set_value("Project Plot Detail", doc.plot, "status", 'VACANT')
	frappe.db.set_value("Project Plot Detail", doc.plot, "sales_rep", None)
	frappe.db.set_value("Project Plot Detail", doc.plot, "offer", None)
	frappe.db.set_value("Project Plot Detail", doc.plot, "date_owned", None)


def update_plots_on_offer_update(doc, method=None):
	if doc.sale_state == "Under Payment" and doc.status == "To Bill":
		frappe.db.set_value("Project Plot Detail", doc.plot, "status", "UNDER PAYMENT")


def update_plots_on_offer_create(doc, method=None):
	project = frappe.get_doc("Project", doc.project)
	for pl in project.plots:
		if pl.name == doc.plot:
			if pl.status != "VACANT":
				frappe.throw("The plot in the offer is NOT VACANT")

	frappe.db.set_value("Project Plot Detail", doc.plot, "customer", doc.customer)
	frappe.db.set_value("Project Plot Detail", doc.plot, "status", 'ON HOLD')
	frappe.db.set_value("Project Plot Detail", doc.plot, "sales_rep", doc.sales_rep)
	frappe.db.set_value("Project Plot Detail", doc.plot, "offer", doc.name)
	frappe.db.set_value("Project Plot Detail", doc.plot, "date_owned", frappe.utils.today())

	doc.agreement = ""


def update_plots_on_sale_submit(doc, method=None):
	project = frappe.get_doc("Project", doc.project)
	for pl in project.plots:
		if pl.name == doc.plot:
			if pl.status != "UNDER PAYMENT":
				frappe.throw("The plot in the offer is UNDER PAYMENT")

	frappe.db.set_value("Project Plot Detail", doc.plot, "status", 'PAID')
	frappe.db.set_value("Project Plot Detail", doc.plot, "sale", doc.name)


@frappe.whitelist()
def get_offers(search_term=None, page_length=None, sort_order=None, start=None, customer=None):
	offer_query = OfferQuery()

	# TODO: check the user and update query for sales person and sales supervisor.
	return offer_query.query(search_term=search_term, page_length=page_length, sort_order=sort_order, start=start, customer=customer)

@frappe.whitelist()
def get_sales(search_term=None, page_length=None, sort_order=None, start=None, customer=None):
	offer_query = SalesQuery()

	# TODO: check the user and update query for sales person and sales supervisor.
	return offer_query.query(search_term=search_term, page_length=page_length, sort_order=sort_order, start=start, customer=customer)

@frappe.whitelist()
def get_payments(search_term=None, page_length=None, sort_order=None, start=None, sale_name=None, offer_name = None):
	offer_query = PaymentsQuery()

	# TODO: check the user and update query for sales person and sales supervisor.
	return offer_query.query(search_term=search_term, page_length=page_length, sort_order=sort_order, start=start, sale_name=sale_name, offer_name=offer_name)

@frappe.whitelist()
def get_payment_methods():
	#TODO: change to get_list for permission
	return frappe.get_all("Mode of Payment", filters=[["enabled", "=", 1]], fields=["name"])


@frappe.whitelist()
def confirm_sale(sale):
	sale = frappe.get_doc("Sales Invoice", sale)
	sale.set_advances()
	sale.submit()
	sale.save()


@frappe.whitelist()
def confirm_offer(date, reference, method, amount, offer_name):
	offer = frappe.get_doc("Sales Order", offer_name)
	#offer.custom_booking_fee = float(amount)
	offer.save()
	offer.submit()
	payment_entry_dict = {
		"company" : offer.company,
		"branch": offer.branch,
		"project": offer.project,
		"payment_type" : "Receive",
		"reference_no" :  reference,
		"reference_date" :  date,
		"party_type" :  "Customer",
		"party" :  offer.customer,
		"posting_date" :  frappe.utils.today(),
		"paid_amount": float(amount),
		"received_amount": float(amount)
	}

	payment_entry = frappe.new_doc("Payment Entry")
	payment_entry.update(payment_entry_dict)
	payment_entry.mode_of_payment = method
	payment_mode = frappe.get_doc("Mode of Payment", method)
	payment_entry.append("references", {"reference_doctype": "Sales Order", 
			"reference_name": offer.name, "allocated_amount": float(amount) })
	
	for account in payment_mode.accounts:
		if account.company == offer.company:
			payment_entry.paid_to = account.default_account
	
	payment_entry.validate()
	payment_entry.insert()

	return offer.name

@frappe.whitelist()
def regenerate_agreement(doc_name):
	doc = frappe.get_doc("Sales Order", doc_name)
	del doc.agreement  

@frappe.whitelist()
def convert_offer(offer):
	offer_doc = frappe.get_doc("Sales Order", offer)
	# sales_inv = frappe.new_doc("Sales Invoice")

	# make_sales_invoice(source_name = offer, target_doc=sales_inv, ignore_permissions=True)

	# sales_inv.payment_schedule = []

	# for off_sch in offer_doc.payment_schedule:
	#     sch = sales_inv.append("payment_schedule", {})
	#     sch.payment_amount = off_sch.payment_amount
	#     sch.due_date = off_sch.due_date
	#     sch.is_deposit = off_sch.is_deposit

	# sales_inv.sales_rep = offer_doc.sales_rep
	# sales_inv.plot = offer_doc.plot
	# sales_inv.set_advances()
	# sales_inv.insert()
	# plot = frappe.get_doc("Plot", offer_doc.plot)
	# plot.status = "UNDER PAYMENT"
	offer_doc.sale_state = "Waiting Sale Approval"
	offer_doc.save()


	return offer_doc.name

@frappe.whitelist()
def update_terms(payment_terms, sale_name):
	sale = frappe.get_doc("Sales Invoice", sale_name)

	sale.payment_schedule = []

	for term in payment_terms:
		sch = sale.append("payment_schedule", {})
		sch.payment_amount = float(term["amount"])
		sch.due_date = term["due_date"]

	sale.save(ignore_permissions=True)

@frappe.whitelist()
def update_offer_terms(terms, offer_name):
	offer = frappe.get_doc("Sales Order", offer_name)

	#offer.payment_schedule = []
	offer.discount_amount = float(terms["discount"])
	
	offer.custom_booking_fee = float(terms["custom_booking_fee"])
	offer.custom_purchase_method = terms["custom_purchase_method"]

	if float(terms["total"]) < 1000000:
		offer.custom_discount_approved = float(terms["discount"]) <= 50000
	
	if float(terms['total']) > 1000000:
		offer.custom_discount_approved = float(terms["discount"]) <= 100000

	for item in offer.items:
		item.rate = float(terms["total"])

	# for term in terms["payment_terms"]:
	#     sch = offer.append("payment_schedule", {})
	#     sch.payment_amount = float(term["amount"])
	#     sch.due_date = term["due_date"]
	#     sch.is_deposit = 1 if term["deposit"] else 0

	del offer.agreement
	offer.save(ignore_permissions=True)


	return "Successful Terms Update"


@frappe.whitelist()
def update_sale_terms_and_convert(terms, offer_name):
	offer = frappe.get_doc("Sales Order", offer_name)

	#offer.payment_schedule = []
	offer.discount_amount = float(terms["discount"])
	offer.custom_booking_fee = float(terms["custom_booking_fee"])
#	offer.custom_purchase_method = terms["custom_purchase_method"]

	if float(terms["total"]) < 1000000:
		offer.custom_discount_approved = float(terms["discount"]) <= 50000
	
	if float(terms['total']) > 1000000:
		offer.custom_discount_approved = float(terms["discount"]) <= 100000

	for item in offer.items:
		item.rate = float(terms["total"])

	offer.payment_schedule = []

	total_is_deposited = 0;

	for term in terms["payment_terms"]:
		sch = offer.append("payment_schedule", {})
		sch.payment_amount = float(term["amount"])
		sch.due_date = term["due_date"]
		sch.is_deposit = 1 if term["deposit"] else 0
		sch.is_deposited = 1 if term["deposited"] else 0

		if sch.is_deposited == 1:
			total_is_deposited += sch.payment_amount

	if total_is_deposited > offer.advance_paid - offer.custom_booking_fee:
		frappe.throw("Payment marked as total deposited is greater than the actual paid for the sale.")


	del offer.agreement
	offer.sale_state = "Waiting Managers Approval"
	offer.save(ignore_permissions=True)

	total = offer.total - offer.discount_amount
	frappe.db.set_value("Sales Order", offer.name, "net_total", total)
	frappe.db.set_value("Sales Order", offer.name, "grand_total", total)
	frappe.db.set_value("Sales Order", offer.name, "rounded_total", total)

	return "Successful Sale Conversion"
 

@frappe.whitelist()
def edit_payment(date, reference, method, amount, payment): 
	payment = frappe.get_doc("Payment Entry", payment)

	if payment.posting_date != date:
		payment.posting_date = date
	
	if payment.reference_no != reference:
		payment.reference_no = reference
	
	if payment.mode_of_payment != method:
		payment.mode_of_payment = method

	if payment.paid_amount != float(amount):
		payment.paid_amount = float(amount)
		payment.received_amount = float(amount)
	payment.validate()
	payment.save()

@frappe.whitelist()
def create_payment(date, reference, method, amount, sale):
	offer = frappe.get_doc("Sales Order", sale)

	payment_entry_dict = {
		"company" : offer.company,
		"branch": offer.branch,
		"project": offer.project,
		"payment_type" : "Receive",
		"reference_no" :  reference,
		"reference_date" :  frappe.utils.today(),
		"party_type" :  "Customer",
		"party" :  offer.customer,
		"posting_date" :  date,
		"paid_amount": float(amount),
		"received_amount": float(amount)
	}
	payment_entry = frappe.new_doc("Payment Entry")
	payment_entry.update(payment_entry_dict)
	payment_entry.mode_of_payment = method
	payment_mode = frappe.get_doc("Mode of Payment", method)
	payment_entry.append("references", {"reference_doctype": "Sales Order", 
			"reference_name": offer.name, "allocated_amount": float(amount) })
	
	for account in payment_mode.accounts:
		if account.company == offer.company:
			payment_entry.paid_to = account.default_account
	
	payment_entry.validate()
	payment_entry.insert()


@frappe.whitelist()
def upload_file(**kwargs):

	files = frappe.request.files
	
	if "file" in files:
		file = files["file"]
		sale = frappe.get_doc("Sales Order", kwargs["sale"])
		filename = kwargs["filename"]

		upload = sale.append("files", {})

		file_doc = frappe.get_doc(
			{
				"doctype": "File",
				"attached_to_doctype": "Sales Order",
				"attached_to_name": sale.name,
				"folder": "Home",
				"file_name": file.filename,
				"is_private": 1,
				"content": file.stream.read(),
			}
		).save(ignore_permissions=True)

		upload.save_name = file_doc.file_name 
		upload.filename = filename
		upload.file = file_doc.file_url

		if filename == "Signed Offer Letter" and sale.sale_state == 'Offer':
			sale.sale_state = "Waiting Signed Offer Letter Approval"

		
		if filename == "Signed Agreement" and sale.sale_state == 'Waiting Agreement Upload':
			sale.sale_state = "Waiting Sale Agreement Approval"

		sale.save(ignore_permissions=True)

	return "Files saved successfully"

@frappe.whitelist()
def get_agreement_templates():
	return frappe.get_all("Print Format", filters={"doc_type": "Sales Order"}, fields=["name"])

@frappe.whitelist()
def update_template_settings(sale, template, include_lawyer):
	sale_doc = frappe.get_doc("Sales Order", sale)
	sale_doc.agreement_template = template
	sale_doc.include_lawyer = include_lawyer
	del sale_doc.agreement
	sale_doc.save(ignore_permissions=True)

@frappe.whitelist()
def download_agreement(doc_name):
	doc = frappe.get_doc("Sales Order", doc_name)
	if doc.agreement is not None and doc.agreement != '':

		generator = PrintFormatGenerator(doc.agreement_template, doc, "NO letter head")
		pdf_file = generator.render_pdf_with_main_html(doc.agreement)

		frappe.local.response.filename = "{name}.pdf".format(
			name=doc_name.replace(" ", "-").replace("/", "-")
		)
		frappe.local.response.filecontent = pdf_file
		frappe.local.response.type = "pdf"
	else:
		download_pdf("Sales Order", doc_name, doc.agreement_template, letterhead=None)


@frappe.whitelist()
def get_sale_cancel_reasons():
	return frappe.get_all("Sale Cancel Reason", fields=["name"])

@frappe.whitelist()
def cancel_sale_request(sale, reason, description=None):

	sale_doc = frappe.get_doc("Sales Order", sale)

	req = frappe.get_doc({
		"doctype": "Sale Update Request", 
		"source": sale,
		"type": "Sale Cancel",
		"cancel_reason": reason,
		"cancel_description": description
	})

	req.insert(ignore_permissions=True)

	sale_doc.sale_state = "Cancel Sale Request"
	sale_doc.save()

	return "Sale Cancel Requested Successfully"


@frappe.whitelist()
def create_transfer_request(sale, project, plot):

	sale_doc = frappe.get_doc("Sales Order", sale)

	sales = frappe.new_doc("Sales Order")
	sales.customer = sale_doc.customer
	sales.sales_rep = sale_doc.sales_rep
	sales.project = project
	sales.branch = sale_doc.branch
	sales.plot = plot
	item = sales.append("items",{})
	item.item_code = 'Plot'
	item.item_name = plot
	item.qty = 1
	item.rate = sale_doc.items[0].rate
	sales.transaction_date = sale_doc.transaction_date
	sales.set_missing_values()

	sales.insert()

	req = frappe.get_doc({
		"doctype": "Sale Update Request", 
		"source": sale,
		"output": sales.name,
		"type": "Plot Transfer",
		"to_project": project,
		"to_plot": plot
	})

	req.insert(ignore_permissions=True)

	sale_doc.sale_state = "Plot Transfer Request"
	sale_doc.save()
	
	sales.sale_state = "Waiting Transfer Approval"
	sales.save()

	return sales.name

@frappe.whitelist()
def delete_offer(offer_name):
    offer = frappe.get_doc("Sales Order", offer_name)

    if offer.sale_state != "Draft":
        frappe.throw("You can't delete this offer because it is not in Draft state.")

    if offer.project and offer.plot:
        project = frappe.get_doc("Project", offer.project)
        
        if not hasattr(project, "project_plot_detail"):
            frappe.throw(f"Project '{project.name}' has no plot details.")

        for ppd in project.project_plot_detail:
            if ppd.plot == offer.plot:
                ppd.status = "VACANT"
                ppd.customer = None
                ppd.date_owned = None
                ppd.sales_rep = None
                ppd.offer = None

        project.save(ignore_permissions=True)
        frappe.db.commit()

    offer.delete()
    frappe.db.commit()

    return "Offer deleted successfully"

class PrintFormatGenerator:
	"""
	Generate a PDF of a Document, with repeatable header and footer if letterhead is provided.

	This generator draws its inspiration and, also a bit of its implementation, from this
	discussion in the library github issues: https://github.com/Kozea/WeasyPrint/issues/92
	"""

	def __init__(self, print_format, doc, letterhead=None):
		"""
		Parameters
		----------
		print_format: str
				Name of the Print Format
		doc: str
				Document to print
		letterhead: str
		----------
				Letter Head to apply (optional)
		"""
		self.base_url = frappe.utils.get_url()
		self.print_format = frappe.get_doc("Print Format", print_format)
		self.doc = doc

		if letterhead == _("No Letterhead"):
			letterhead = None
		self.letterhead = frappe.get_doc("Letter Head", letterhead) if letterhead else None

		self.build_context()
		self.layout = self.get_layout(self.print_format)
		self.context.layout = self.layout

	def build_context(self):
		self.print_settings = frappe.get_doc("Print Settings")
		page_width_map = {"A4": 210, "Letter": 216}
		page_width = page_width_map.get(self.print_settings.pdf_page_size) or 210
		body_width = page_width - self.print_format.margin_left - self.print_format.margin_right
		print_style = (
			frappe.get_doc("Print Style", self.print_settings.print_style)
			if self.print_settings.print_style
			else None
		)
		context = frappe._dict(
			{
				"doc": self.doc,
				"print_format": self.print_format,
				"print_settings": self.print_settings,
				"print_style": print_style,
				"letterhead": self.letterhead,
				"page_width": page_width,
				"body_width": body_width,
			}
		)
		self.context = context

	def get_html_preview(self):
		header_html, footer_html = self.get_header_footer_html()
		self.context.header = header_html
		self.context.footer = footer_html
		return self.get_main_html()

	def get_main_html(self):
		self.context.css = frappe.render_template("templates/print_format/print_format.css", self.context)
		return frappe.render_template("templates/print_format/print_format.html", self.context)

	def get_header_footer_html(self):
		header_html = footer_html = None
		if self.letterhead:
			header_html = frappe.render_template("templates/print_format/print_header.html", self.context)
		if self.letterhead:
			footer_html = frappe.render_template("templates/print_format/print_footer.html", self.context)
		return header_html, footer_html

	def render_pdf(self):
		"""
		Returns
		-------
		pdf: a bytes sequence
				The rendered PDF.
		"""
		HTML, CSS = import_weasyprint()

		self._make_header_footer()

		self.context.update({"header_height": self.header_height, "footer_height": self.footer_height})
		main_html = self.get_main_html()

		html = HTML(string=main_html, base_url=self.base_url)
		main_doc = html.render()

		if self.header_html or self.footer_html:
			self._apply_overlay_on_main(main_doc, self.header_body, self.footer_body)
		return main_doc.write_pdf()

	def render_pdf_without_headers(self):
		"""
		Returns
		-------
		pdf: a bytes sequence
				The rendered PDF.
		"""
		HTML, CSS = import_weasyprint()

		self._make_header_footer()

		self.context.update({"header_height": self.header_height, "footer_height": self.footer_height})
		main_html = self.get_main_html()
		
		main_html = main_html.replace('content: counter(page) " of " counter(pages)', "")

		html = HTML(string=main_html, base_url=self.base_url)
		main_doc = html.render()
		
		return main_doc.write_pdf()
	
	def render_pdf_with_main_html(self, main_html):
		"""
		Returns
		-------
		pdf: a bytes sequence
				The rendered PDF.
		"""
		HTML, CSS = import_weasyprint()

		self._make_header_footer()

		self.context.update({"header_height": self.header_height, "footer_height": self.footer_height})


		html = HTML(string=main_html, base_url=self.base_url)
		main_doc = html.render()

		if self.header_html or self.footer_html:
			self._apply_overlay_on_main(main_doc, self.header_body, self.footer_body)
		return main_doc.write_pdf()


	def render_pdf_with_main_html_without_headers(self, main_html):
		"""
		Returns
		-------
		pdf: a bytes sequence
				The rendered PDF.
		"""
		HTML, CSS = import_weasyprint()

		self._make_header_footer()

		self.context.update({"header_height": self.header_height, "footer_height": self.footer_height})

		main_html = main_html.replace('content: counter(page) " of " counter(pages)', "")

		html = HTML(string=main_html, base_url=self.base_url)
		main_doc = html.render()
		return main_doc.write_pdf()

		


	def _compute_overlay_element(self, element: str):
		"""
		Parameters
		----------
		element: str
				Either 'header' or 'footer'

		Returns
		-------
		element_body: BlockBox
				A Weasyprint pre-rendered representation of an html element
		element_height: float
				The height of this element, which will be then translated in a html height
		"""
		HTML, CSS = import_weasyprint()

		html = HTML(
			string=getattr(self, f"{element}_html"),
			base_url=self.base_url,
		)
		element_doc = html.render(stylesheets=[CSS(string="@page {size: A4 portrait; margin: 0;}")])
		element_page = element_doc.pages[0]
		element_body = PrintFormatGenerator.get_element(element_page._page_box.all_children(), "body")
		element_body = element_body.copy_with_children(element_body.all_children())
		element_html = PrintFormatGenerator.get_element(element_page._page_box.all_children(), element)

		if element == "header":
			element_height = element_html.height
		if element == "footer":
			element_height = element_page.height - element_html.position_y

		return element_body, element_height

	def _apply_overlay_on_main(self, main_doc, header_body=None, footer_body=None):
		"""
		Insert the header and the footer in the main document.

		Parameters
		----------
		main_doc: Document
				The top level representation for a PDF page in Weasyprint.
		header_body: BlockBox
				A representation for an html element in Weasyprint.
		footer_body: BlockBox
				A representation for an html element in Weasyprint.
		"""
		for page in main_doc.pages:
			page_body = PrintFormatGenerator.get_element(page._page_box.all_children(), "body")

			if header_body:
				page_body.children += header_body.all_children()
			if footer_body:
				page_body.children += footer_body.all_children()

	def _make_header_footer(self):
		self.header_html, self.footer_html = self.get_header_footer_html()

		if self.header_html:
			header_body, header_height = self._compute_overlay_element("header")
		else:
			header_body, header_height = None, 0
		if self.footer_html:
			footer_body, footer_height = self._compute_overlay_element("footer")
		else:
			footer_body, footer_height = None, 0

		self.header_body = header_body
		self.header_height = header_height
		self.footer_body = footer_body
		self.footer_height = footer_height

	def get_layout(self, print_format):
		layout = frappe.parse_json(print_format.format_data)
		layout = self.set_field_renderers(layout)
		layout = self.process_margin_texts(layout)
		return layout

	def set_field_renderers(self, layout):
		renderers = {"HTML Editor": "HTML", "Markdown Editor": "Markdown"}
		for section in layout["sections"]:
			for column in section["columns"]:
				for df in column["fields"]:
					fieldtype = df["fieldtype"]
					renderer_name = fieldtype.replace(" ", "")
					df["renderer"] = renderers.get(fieldtype) or renderer_name
					df["section"] = section
		return layout

	def process_margin_texts(self, layout):
		margin_texts = [
			"top_left",
			"top_center",
			"top_right",
			"bottom_left",
			"bottom_center",
			"bottom_right",
		]
		for key in margin_texts:
			text = layout.get("text_" + key)
			if text and "{{" in text:
				layout["text_" + key] = frappe.render_template(text, self.context)

		return layout

	@staticmethod
	def get_element(boxes, element):
		"""
		Given a set of boxes representing the elements of a PDF page in a DOM-like way, find the
		box which is named `element`.

		Look at the notes of the class for more details on Weasyprint insides.
		"""
		for box in boxes:
			if box.element_tag == element:
				return box
			return PrintFormatGenerator.get_element(box.all_children(), element)


def import_weasyprint():
	try:
		from weasyprint import CSS, HTML

		return HTML, CSS
	except OSError:
		message = "\n".join(
			[
				"WeasyPrint depdends on additional system dependencies.",
				"Follow instructions specific to your operating system:",
				"https://doc.courtbouillon.org/weasyprint/stable/first_steps.html",
			]
		)
		click.secho(message, fg="yellow")
		frappe.throw(message)

@frappe.whitelist()
def download_word_document(doc_name):

	doc = frappe.get_doc("Sales Order", doc_name)
	

	generator = PrintFormatGenerator(doc.agreement_template, doc, "NO letter head")
	if doc.agreement is not None and doc.agreement != '':

		pdf_file = generator.render_pdf_with_main_html_without_headers(doc.agreement)
	else: 
		pdf_file = generator.render_pdf_without_headers()

	
	# Save PDF to a temporary file
	with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
		temp_pdf.write(pdf_file)  # Write PDF bytes
		temp_pdf_path = temp_pdf.name  

	word_stream = BytesIO()
	
	cv = Converter(temp_pdf_path)  # Create Converter instance
	cv.convert(word_stream)  # Convert PDF to DOCX
	cv.close() 

	word_stream.seek(0)

	document = Document(word_stream)

	# Add header and footer
	add_header_footer(document, doc)

	# Save the updated document back to a new BytesIO stream
	updated_word_stream = BytesIO()
	document.save(updated_word_stream)
	updated_word_stream.seek(0)


	# Prepare response
	frappe.local.response.filename = f"{doc_name.replace(' ', '-').replace('/', '-')}.docx"
	frappe.local.response.filecontent = updated_word_stream.getvalue()
	frappe.local.response.type = "binary"


def add_header_footer(doc, sale_doc):
    """Adds a header and 'Page X of Y' to the footer of a Word document."""

    # 🔹 Get the first section of the document
    section = doc.sections[0]

    # 🔹 Add a header
    header = section.header
    header_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_paragraph.text = "Sales Order Agreement"
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center the header
    header_paragraph.style.font.size = Pt(12)

    # 🔹 Add a footer
    footer = section.footer
    footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center the footer

    # 🔹 Insert "Page "
    footer_paragraph.add_run("Page ")

    # 🔹 Insert field for current page number
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText1 = OxmlElement('w:instrText')
    instrText1.text = "PAGE"
    instrText1.set(qn('xml:space'), 'preserve')

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run = footer_paragraph.add_run()
    run._r.append(fldChar1)
    run._r.append(instrText1)
    run._r.append(fldChar2)

    # 🔹 Insert " of "
    footer_paragraph.add_run(" of ")

    # 🔹 Insert field for total number of pages
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')

    instrText2 = OxmlElement('w:instrText')
    instrText2.text = "NUMPAGES"
    instrText2.set(qn('xml:space'), 'preserve')

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    run2 = footer_paragraph.add_run()
    run2._r.append(fldChar3)
    run2._r.append(instrText2)
    run2._r.append(fldChar4)
